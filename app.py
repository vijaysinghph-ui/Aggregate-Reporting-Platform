import io
import zipfile
import xml.etree.ElementTree as ET
from datetime import date

import pandas as pd
import streamlit as st
from docx import Document
from openai import OpenAI
from pypdf import PdfReader

# =========================================================
# App Config
# =========================================================
st.set_page_config(
    page_title="Viginovix Aggregate Reporting Platform",
    layout="wide"
)

# =========================================================
# Constants
# =========================================================
REPORT_TYPES = {
    "PADER": {
        "sections": [
            {
                "id": "cover_page",
                "title": "Cover Page",
                "purpose": "Capture title page details such as report title, product, strength, NDA/ANDA number, reporting period, company details, confidentiality statement, approval date, report status, and date of report."
            },
            {
                "id": "approval_page",
                "title": "Approval Page",
                "purpose": "Capture author, medical reviewer, reviewer, and approver details with signature/date placeholders."
            },
            {
                "id": "table_of_contents",
                "title": "Table of Contents",
                "purpose": "Auto-generate the table of contents for the assembled report."
            },
            {
                "id": "introduction",
                "title": "1. Introduction",
                "purpose": "Draft the Introduction using current report setup, previous PADER reference if available, and current label information if available."
            },
            {
                "id": "summary_alerts_new_ades_followup",
                "title": "2. Summary of Submitted 15-Day Alerts, New Adverse Drug Experiences and New Adverse Drug Experience Follow-up",
                "purpose": "Draft the core safety summary using uploaded line listing data, with backend logic determining the best structure."
            },
            {
                "id": "actions_taken",
                "title": "3. Actions Taken Since Last Periodic Adverse Drug Experience Report",
                "purpose": "Summarize product-specific regulatory actions, labeling changes, safety actions, and related authority actions during the reporting period."
            },
            {
                "id": "conclusion",
                "title": "4. Conclusion",
                "purpose": "Provide the overall safety conclusion and state whether the product safety profile remains unchanged or if further action is planned."
            }
        ]
    },
    "PBRER": {"sections": []},
    "DSUR": {"sections": []}
}

# =========================================================
# OpenAI Helper
# =========================================================
def get_openai_client():
    api_key = st.secrets.get("OPENAI_API_KEY", None)
    if not api_key:
        return None
    return OpenAI(api_key=api_key)

def safe_text(value) -> str:
    return "" if value is None else str(value)

# =========================================================
# File Extraction Helpers
# =========================================================
def extract_text_from_pdf(uploaded_file) -> str:
    """
    Extract text from searchable PDFs.
    If the PDF is image-only/scanned, output may be poor or empty.
    """
    try:
        uploaded_file.seek(0)
        reader = PdfReader(uploaded_file)
        texts = []

        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                texts.append(page_text)

        return "\n".join(texts).strip()
    except Exception:
        return ""

def extract_text_from_docx(uploaded_file) -> str:
    """
    Basic DOCX text extraction without extra packages.
    """
    try:
        uploaded_file.seek(0)
        with zipfile.ZipFile(uploaded_file) as z:
            xml_content = z.read("word/document.xml")
        root = ET.fromstring(xml_content)
        texts = []
        for node in root.iter():
            if node.tag.endswith("}t") and node.text:
                texts.append(node.text)
        return "\n".join(texts).strip()
    except Exception:
        return ""

def extract_text_from_txt(uploaded_file) -> str:
    try:
        uploaded_file.seek(0)
        return uploaded_file.read().decode("utf-8", errors="ignore")
    except Exception:
        return ""

def extract_reference_text(uploaded_file) -> str:
    """
    Supports PDF, DOCX, and TXT.
    """
    if uploaded_file is None:
        return ""

    filename = uploaded_file.name.lower()

    if filename.endswith(".pdf"):
        return extract_text_from_pdf(uploaded_file)

    if filename.endswith(".docx"):
        return extract_text_from_docx(uploaded_file)

    if filename.endswith(".txt"):
        return extract_text_from_txt(uploaded_file)

    return ""

# =========================================================
# Table Helpers
# =========================================================
def read_uploaded_table(uploaded_file):
    if uploaded_file is None:
        return None
    try:
        if uploaded_file.name.lower().endswith(".csv"):
            return pd.read_csv(uploaded_file)
        return pd.read_excel(uploaded_file)
    except Exception as e:
        st.error(f"Error reading file: {str(e)}")
        return None

def detect_column(df: pd.DataFrame, candidates: list[str]):
    cols = list(df.columns)
    lower_map = {str(c).strip().lower(): c for c in cols}

    for cand in candidates:
        cand_l = cand.strip().lower()
        if cand_l in lower_map:
            return lower_map[cand_l]

    for cand in candidates:
        cand_l = cand.strip().lower()
        for c in cols:
            if cand_l in str(c).strip().lower():
                return c

    return None

def summarize_dataframe(df: pd.DataFrame, max_rows: int = 10) -> str:
    if df is None or df.empty:
        return "No data available."

    lines = [f"Total rows: {len(df)}", f"Columns: {list(df.columns)}", "", "Sample rows:"]
    preview = df.head(max_rows)
    for _, row in preview.iterrows():
        row_text = " | ".join([f"{col}: {row[col]}" for col in preview.columns[:8]])
        lines.append(f"- {row_text}")
    return "\n".join(lines)

def build_line_listing_backend_summary(df: pd.DataFrame) -> str:
    """
    Backend-oriented Section 2 summary.
    """
    if df is None or df.empty:
        return "No line listing data available."

    case_id_col = detect_column(df, ["case id", "case number", "case no", "case_id"])
    event_col = detect_column(df, ["event term", "adverse drug experiences", "event", "pt", "preferred term"])
    soc_col = detect_column(df, ["soc", "system organ class"])
    seriousness_col = detect_column(df, ["seriousness", "serious"])
    listedness_col = detect_column(df, ["listedness", "listed/unlisted", "listedness status"])
    causality_col = detect_column(df, ["causality", "relatedness", "causal association"])
    report_type_col = detect_column(df, ["report type"])
    expedited_col = detect_column(df, ["expedited status", "expedited"])
    followup_col = detect_column(df, ["follow-up", "follow up"])
    outcome_col = detect_column(df, ["outcome"])
    country_col = detect_column(df, ["country"])

    lines = []
    lines.append(f"Total number of cases: {len(df)}")
    lines.append("Detected columns:")
    lines.append(f"- Case ID: {case_id_col}")
    lines.append(f"- Event Term: {event_col}")
    lines.append(f"- SOC: {soc_col}")
    lines.append(f"- Seriousness: {seriousness_col}")
    lines.append(f"- Listedness: {listedness_col}")
    lines.append(f"- Causality: {causality_col}")
    lines.append(f"- Report Type: {report_type_col}")
    lines.append(f"- Expedited Status: {expedited_col}")
    lines.append(f"- Follow-up: {followup_col}")
    lines.append(f"- Outcome: {outcome_col}")
    lines.append(f"- Country: {country_col}")

    expedited_df = None
    if expedited_col:
        expedited_df = df[df[expedited_col].astype(str).str.lower().isin(["yes", "y", "true", "1", "expedited"])]
    elif report_type_col:
        expedited_df = df[df[report_type_col].astype(str).str.lower().str.contains("15|alert|expedited", na=False)]

    followup_df = None
    if followup_col:
        followup_df = df[df[followup_col].astype(str).str.lower().isin(["yes", "y", "true", "1", "follow-up", "follow up"])]
    elif report_type_col:
        followup_df = df[df[report_type_col].astype(str).str.lower().str.contains("follow", na=False)]

    serious_df = None
    if seriousness_col:
        serious_df = df[df[seriousness_col].astype(str).str.lower().str.contains("serious|yes|y|true", na=False)]

    unlisted_df = None
    if listedness_col:
        unlisted_df = df[df[listedness_col].astype(str).str.lower().str.contains("unlisted", na=False)]

    related_df = None
    if causality_col:
        related_df = df[df[causality_col].astype(str).str.lower().str.contains("related|causal", na=False)]

    sur_df = df.copy()
    if seriousness_col:
        sur_df = sur_df[sur_df[seriousness_col].astype(str).str.lower().str.contains("serious|yes|y|true", na=False)]
    if listedness_col:
        sur_df = sur_df[sur_df[listedness_col].astype(str).str.lower().str.contains("unlisted", na=False)]
    if causality_col:
        sur_df = sur_df[sur_df[causality_col].astype(str).str.lower().str.contains("related|causal", na=False)]

    lines.append("")
    lines.append("Core classification summary:")
    lines.append(f"- Expedited / 15-day alert cases: {len(expedited_df) if expedited_df is not None else 'Not determined'}")
    lines.append(f"- Follow-up cases: {len(followup_df) if followup_df is not None else 'Not determined'}")
    lines.append(f"- Serious cases: {len(serious_df) if serious_df is not None else 'Not determined'}")
    lines.append(f"- Unlisted cases: {len(unlisted_df) if unlisted_df is not None else 'Not determined'}")
    lines.append(f"- Related cases: {len(related_df) if related_df is not None else 'Not determined'}")
    lines.append(f"- Serious unlisted related cases: {len(sur_df)}")

    if outcome_col:
        fatal_df = df[df[outcome_col].astype(str).str.lower().str.contains("fatal|death|died", na=False)]
        lines.append(f"- Fatal cases: {len(fatal_df)}")

    if soc_col:
        lines.append("")
        lines.append("Top SOC distribution:")
        soc_counts = df[soc_col].astype(str).value_counts(dropna=False).head(10)
        for idx, val in soc_counts.items():
            lines.append(f"- {idx}: {val}")

    if event_col:
        lines.append("")
        lines.append("Top event terms:")
        event_counts = df[event_col].astype(str).value_counts(dropna=False).head(10)
        for idx, val in event_counts.items():
            lines.append(f"- {idx}: {val}")

    if len(sur_df) > 0:
        lines.append("")
        lines.append("Serious unlisted related case examples:")
        cols_to_show = [c for c in [case_id_col, event_col, soc_col, causality_col, listedness_col, outcome_col] if c]
        sample_sur = sur_df[cols_to_show].head(8) if cols_to_show else sur_df.head(5)
        for _, row in sample_sur.iterrows():
            row_text = " | ".join([f"{col}: {row[col]}" for col in sample_sur.columns])
            lines.append(f"- {row_text}")

    return "\n".join(lines)

# =========================================================
# AI Drafting Helpers
# =========================================================
def generate_ai_draft(
    section_title: str,
    section_purpose: str,
    source_data: str,
    comments: str,
    product_name: str,
    interval_start,
    interval_end,
) -> str:
    client = get_openai_client()
    if client is None:
        return "ERROR: OPENAI_API_KEY not found in Streamlit secrets."

    if not safe_text(source_data).strip():
        return "ERROR: Please provide source data before generating a draft."

    instructions = (
        "You are an expert pharmacovigilance medical writer. "
        "Draft only the requested report section in a concise, professional, regulatory style. "
        "Use only the source data provided. "
        "Do not invent facts, numbers, dates, tables, or conclusions. "
        "If information is missing, stay neutral and do not hallucinate. "
        "Do not repeat the section heading if it is already provided outside the body text. "
        "Return only the drafted section body text."
    )

    user_input = f"""
Report Type: PADER
Section Title: {section_title}
Section Purpose: {section_purpose}

Report Context:
Product Name: {product_name}
Reporting Interval Start: {interval_start}
Reporting Interval End: {interval_end}

Source Data:
{source_data}

Drafting Instructions:
{comments}
"""

    try:
        response = client.responses.create(
            model="gpt-5.4",
            instructions=instructions,
            input=user_input,
        )
        return response.output_text.strip()
    except Exception as e:
        return f"ERROR: {str(e)}"

def generate_introduction_draft(
    report_context_text: str,
    previous_pader_text: str,
    label_text: str,
) -> str:
    client = get_openai_client()
    if client is None:
        return "ERROR: OPENAI_API_KEY not found in Streamlit secrets."

    instructions = (
        "You are an expert pharmacovigilance medical writer drafting the Introduction section of a PADER. "
        "Use the current report context as authoritative for reporting interval and metadata. "
        "Use current label text as the primary source for current product truth when available. "
        "Use previous PADER text only as historical reference for stable wording and style. "
        "If a previous PADER is not available, generate the Introduction using report context and label only. "
        "If uploaded PDFs produce limited text, rely on the available extracted text and report context without guessing. "
        "Do not hallucinate product facts, regulatory details, or placeholders. "
        "Write gracefully even if some information is missing. "
        "Do not repeat the section title. "
        "Return only the Introduction body text."
    )

    user_input = f"""
Current Report Context:
{report_context_text}

Previous PADER Reference Text:
{previous_pader_text}

Current Label Reference Text:
{label_text}
"""

    try:
        response = client.responses.create(
            model="gpt-5.4",
            instructions=instructions,
            input=user_input,
        )
        return response.output_text.strip()
    except Exception as e:
        return f"ERROR: {str(e)}"

def generate_section2_draft(
    product_name: str,
    interval_start,
    interval_end,
    line_listing_summary: str,
) -> str:
    client = get_openai_client()
    if client is None:
        return "ERROR: OPENAI_API_KEY not found in Streamlit secrets."

    if not safe_text(line_listing_summary).strip():
        return "ERROR: No line listing summary available."

    instructions = (
        "You are an expert pharmacovigilance physician writing the PADER section "
        "'Summary of Submitted 15-Day Alerts, New Adverse Drug Experiences and New Adverse Drug Experience Follow-up'. "
        "Use only the provided backend line listing summary. "
        "Do not invent counts, cases, dates, causality, listedness, seriousness, or conclusions. "
        "At the backend, align the output in a medically useful structure: "
        "first address 15-day alerts if identifiable; "
        "then emphasize medically important serious unlisted related cases if present; "
        "then describe other notable new adverse drug experiences and follow-up information if present; "
        "then end with a restrained concluding statement only if supported by the data. "
        "If fatal events are present, mention them in a neutral, regulatory way. "
        "If SOC patterns are present, you may group discussion accordingly. "
        "If causality information is present, use it carefully and only as reported. "
        "Maintain a cautious regulatory tone similar to real PADER narratives. "
        "Do not repeat the section title. "
        "Return only the section body text."
    )

    user_input = f"""
Product Name: {product_name}
Reporting Interval: {interval_start} to {interval_end}

Backend Line Listing Summary:
{line_listing_summary}
"""

    try:
        response = client.responses.create(
            model="gpt-5.4",
            instructions=instructions,
            input=user_input,
        )
        return response.output_text.strip()
    except Exception as e:
        return f"ERROR: {str(e)}"

def generate_actions_taken_draft(
    regulatory_actions_summary: str,
    product_name: str,
    interval_start,
    interval_end
) -> str:
    client = get_openai_client()
    if client is None:
        return "ERROR: OPENAI_API_KEY not found in Streamlit secrets."

    if not safe_text(regulatory_actions_summary).strip():
        return (
            "No actions related to safety, labeling, or regulatory authority decisions "
            "were identified during the reporting interval."
        )

    instructions = (
        "You are an expert pharmacovigilance medical writer drafting the PADER section "
        "'Actions Taken Since Last Periodic Adverse Drug Experience Report'. "
        "Use only the uploaded regulatory actions summary. "
        "Do not invent actions, approvals, or authority decisions. "
        "If no meaningful actions are evident, state that no actions were identified during the reporting interval. "
        "Do not repeat the section title. "
        "Return only the section body text."
    )

    user_input = f"""
Product Name: {product_name}
Reporting Interval: {interval_start} to {interval_end}

Regulatory Actions Summary:
{regulatory_actions_summary}
"""

    try:
        response = client.responses.create(
            model="gpt-5.4",
            instructions=instructions,
            input=user_input,
        )
        return response.output_text.strip()
    except Exception as e:
        return f"ERROR: {str(e)}"

# =========================================================
# Static Section Builders
# =========================================================
def generate_cover_page_text(
    product_name: str,
    dosage_strength: str,
    nda_anda_number: str,
    company_name: str,
    interval_start,
    interval_end,
    approval_date,
    report_status: str,
    report_status_other: str,
    report_date,
    confidentiality_statement: str,
    company_address: str
) -> str:
    final_status = report_status_other if report_status == "Other" and report_status_other else report_status
    lines = [
        "ANNUAL ADVERSE DRUG EXPERIENCE REPORT",
        f"({interval_start} to {interval_end})",
        f"{product_name}, {dosage_strength}; NDA/ANDA No. {nda_anda_number}",
        "",
        f"{company_name}",
        company_address,
        "",
        "Periodic Adverse Drug Experience Report",
        "A report for the United States Food and Drug Administration",
        "",
        f"Approval Date: {approval_date}",
        f"Review Period: {interval_start} to {interval_end}",
        f"Report Status: {final_status}",
        f"Date of Report: {report_date}",
        "",
        confidentiality_statement
    ]
    return "\n".join(lines)

def generate_approval_page_text(
    product_name: str,
    interval_start,
    interval_end,
    author_name: str,
    author_designation: str,
    medical_reviewer_name: str,
    medical_reviewer_designation: str,
    reviewer_name: str,
    reviewer_designation: str,
    approver_name: str,
    approver_designation: str,
    company_name: str
) -> str:
    lines = [
        f"Product: {product_name}",
        f"Reporting Interval: {interval_start} to {interval_end}",
        "",
        "Author:",
        f"Name: {author_name}",
        f"Designation: {author_designation}",
        f"For: {company_name}",
        "Signature: ____________________",
        "Date: ____________________",
        "",
        "Medical Reviewer:",
        f"Name: {medical_reviewer_name}",
        f"Designation: {medical_reviewer_designation}",
        f"For: {company_name}",
        "Signature: ____________________",
        "Date: ____________________",
        "",
        "Reviewed by:",
        f"Name: {reviewer_name}",
        f"Designation: {reviewer_designation}",
        f"For: {company_name}",
        "Signature: ____________________",
        "Date: ____________________",
        "",
        "Approved by:",
        f"Name: {approver_name}",
        f"Designation: {approver_designation}",
        f"For: {company_name}",
        "Signature: ____________________",
        "Date: ____________________",
    ]
    return "\n".join(lines)

def generate_toc_text(sections: list[dict]) -> str:
    toc_lines = []
    for section in sections:
        if section["id"] == "cover_page":
            continue
        toc_lines.append(section["title"])
    return "\n".join(toc_lines)

# =========================================================
# Assembly + Export
# =========================================================
def assemble_full_report(
    product_name: str,
    interval_start,
    interval_end,
    report_owner: str,
    sections: list[dict]
) -> str:
    report_parts = []

    report_parts.append("ANNUAL ADVERSE DRUG EXPERIENCE REPORT")
    report_parts.append(f"Product: {product_name}")
    report_parts.append(f"Review Period: {interval_start} to {interval_end}")
    report_parts.append(f"Report Owner: {report_owner}")
    report_parts.append("\n" + "=" * 80 + "\n")

    for section in sections:
        section_title = section["title"]
        draft_key = f"draft_{section['id']}"
        section_text = safe_text(st.session_state.get(draft_key, "")).strip()

        if not section_text:
            section_text = "[No draft available for this section yet.]"

        report_parts.append(section_title)
        report_parts.append(section_text)
        report_parts.append("\n" + "-" * 80 + "\n")

    return "\n".join(report_parts)

def export_report_to_word(full_report_text: str) -> bytes:
    doc = Document()
    lines = full_report_text.splitlines()

    for line in lines:
        stripped = line.strip()

        if not stripped:
            doc.add_paragraph("")
        elif stripped == "ANNUAL ADVERSE DRUG EXPERIENCE REPORT":
            doc.add_heading(stripped, level=0)
        elif stripped in [
            "Cover Page",
            "Approval Page",
            "Table of Contents",
            "1. Introduction",
            "2. Summary of Submitted 15-Day Alerts, New Adverse Drug Experiences and New Adverse Drug Experience Follow-up",
            "3. Actions Taken Since Last Periodic Adverse Drug Experience Report",
            "4. Conclusion",
        ]:
            doc.add_heading(stripped, level=1)
        else:
            doc.add_paragraph(line)

    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio.getvalue()

# =========================================================
# Main UI
# =========================================================
st.title("Viginovix Aggregate Reporting Platform")
st.write("Prototype: AI-assisted aggregate report authoring and review")

st.header("Step 1: Select Report Type")

report_type = st.selectbox(
    "Choose Report Type",
    ["Select...", "PADER", "PBRER", "DSUR"]
)

if report_type == "PADER":
    st.success("PADER selected. Sections will load below.")

    # -----------------------------------------------------
    # Step 2: Report Setup
    # -----------------------------------------------------
    st.header("Step 2: Report Setup")

    col1, col2 = st.columns(2)

    with col1:
        product_name = st.text_input("Product Name")
        nda_anda_number = st.text_input("NDA / ANDA Number")
        approval_date = st.date_input("Approval Date", value=date.today())
        company_name = st.text_input("Company Name")
        dosage_strength = st.text_input("Dosage Form / Strength")
        company_address = st.text_area("Company Address", height=100)

    with col2:
        interval_start = st.date_input("Reporting Interval Start Date", value=date.today())
        interval_end = st.date_input("Reporting Interval End Date", value=date.today())
        data_lock_point = st.date_input("Data Lock Point", value=date.today())
        region = st.selectbox("Region", ["US", "EU", "UK", "Global"])
        template_version = st.text_input("Template Version", value="v1.0")
        report_owner = st.text_input("Report Owner")
        report_status = st.selectbox(
            "Report Status",
            ["Annual", "Quarterly", "Other"]
        )

    report_status_other = ""
    if report_status == "Other":
        report_status_other = st.text_input("If Other, specify Report Status")

    report_date = st.date_input("Date of Report", value=date.today())
    confidentiality_statement = st.text_area(
        "Confidentiality Statement",
        value="This document is a confidential communication. Acceptance of this document constitutes an agreement by the recipient that no unpublished information contained herein will be published or disclosed without prior written approval.",
        height=100
    )

    # -----------------------------------------------------
    # Approval workflow fields
    # -----------------------------------------------------
    st.header("Approval Workflow Details")

    a1, a2 = st.columns(2)
    with a1:
        author_name = st.text_input("Author Name")
        author_designation = st.text_input("Author Designation")
        medical_reviewer_name = st.text_input("Medical Reviewer Name")
        medical_reviewer_designation = st.text_input("Medical Reviewer Designation")
    with a2:
        reviewer_name = st.text_input("Reviewer Name")
        reviewer_designation = st.text_input("Reviewer Designation")
        approver_name = st.text_input("Approver Name")
        approver_designation = st.text_input("Approver Designation")

    # -----------------------------------------------------
    # Step 3: PADER Sections
    # -----------------------------------------------------
    st.header("Step 3: PADER Sections")

    pader_sections = REPORT_TYPES["PADER"]["sections"]

    for section in pader_sections:
        with st.expander(section["title"]):
            st.write(f"**Purpose:** {section['purpose']}")

            # Cover Page
            if section["id"] == "cover_page":
                if st.button("Generate Cover Page", key="btn_cover_page"):
                    st.session_state["draft_cover_page"] = generate_cover_page_text(
                        product_name=product_name,
                        dosage_strength=dosage_strength,
                        nda_anda_number=nda_anda_number,
                        company_name=company_name,
                        interval_start=interval_start,
                        interval_end=interval_end,
                        approval_date=approval_date,
                        report_status=report_status,
                        report_status_other=report_status_other,
                        report_date=report_date,
                        confidentiality_statement=confidentiality_statement,
                        company_address=company_address
                    )

                st.text_area(
                    "Draft Output for Cover Page",
                    key="draft_cover_page",
                    height=280
                )

            # Approval Page
            elif section["id"] == "approval_page":
                if st.button("Generate Approval Page", key="btn_approval_page"):
                    st.session_state["draft_approval_page"] = generate_approval_page_text(
                        product_name=product_name,
                        interval_start=interval_start,
                        interval_end=interval_end,
                        author_name=author_name,
                        author_designation=author_designation,
                        medical_reviewer_name=medical_reviewer_name,
                        medical_reviewer_designation=medical_reviewer_designation,
                        reviewer_name=reviewer_name,
                        reviewer_designation=reviewer_designation,
                        approver_name=approver_name,
                        approver_designation=approver_designation,
                        company_name=company_name
                    )

                st.text_area(
                    "Draft Output for Approval Page",
                    key="draft_approval_page",
                    height=320
                )

            # TOC
            elif section["id"] == "table_of_contents":
                if st.button("Generate Table of Contents", key="btn_table_of_contents"):
                    st.session_state["draft_table_of_contents"] = generate_toc_text(pader_sections)

                st.text_area(
                    "Draft Output for Table of Contents",
                    key="draft_table_of_contents",
                    height=220
                )

            # Section 1 simplified
            elif section["id"] == "introduction":
                previous_pader_file = st.file_uploader(
                    "Upload Previous PADER (optional)",
                    type=["pdf", "docx", "txt"],
                    key="previous_pader_upload"
                )

                label_file = st.file_uploader(
                    "Upload Current Label",
                    type=["pdf", "docx", "txt"],
                    key="label_upload"
                )

                if st.button("Generate Introduction", key="btn_introduction"):
                    with st.spinner("Generating Introduction..."):
                        previous_pader_text = extract_reference_text(previous_pader_file)
                        label_text = extract_reference_text(label_file)

                        report_context_text = f"""
Product Name: {product_name}
NDA / ANDA Number: {nda_anda_number}
Approval Date: {approval_date}
Company Name: {company_name}
Dosage Form / Strength: {dosage_strength}
Reporting Interval Start: {interval_start}
Reporting Interval End: {interval_end}
Region: {region}
Report Status: {report_status_other if report_status == 'Other' and report_status_other else report_status}
"""

                        if not previous_pader_text and not label_text:
                            st.warning(
                                "No extractable text was found from the uploaded files. "
                                "If these are scanned PDFs, this prototype may not extract them reliably."
                            )

                        draft_text = generate_introduction_draft(
                            report_context_text=report_context_text,
                            previous_pader_text=previous_pader_text,
                            label_text=label_text,
                        )
                        st.session_state["draft_introduction"] = draft_text

                st.text_area(
                    "Draft Output for 1. Introduction",
                    key="draft_introduction",
                    height=260
                )

            # Section 2 simplified
            elif section["id"] == "summary_alerts_new_ades_followup":
                uploaded_file = st.file_uploader(
                    "Upload PADER Line Listing",
                    type=["xlsx", "csv"],
                    key="line_listing_upload"
                )

                if uploaded_file is not None:
                    df = read_uploaded_table(uploaded_file)
                    if df is not None:
                        st.session_state["line_listing_df"] = df
                        st.success("Line listing uploaded successfully.")
                        st.dataframe(df.head(10), use_container_width=True)

                if st.button("Generate Section 2", key="btn_summary_alerts_new_ades_followup"):
                    df = st.session_state.get("line_listing_df", None)
                    if df is None:
                        st.error("Please upload a line listing file first.")
                    else:
                        with st.spinner("Generating Section 2..."):
                            backend_summary = build_line_listing_backend_summary(df)
                            draft_text = generate_section2_draft(
                                product_name=product_name,
                                interval_start=interval_start,
                                interval_end=interval_end,
                                line_listing_summary=backend_summary,
                            )
                            st.session_state["draft_summary_alerts_new_ades_followup"] = draft_text

                st.text_area(
                    "Draft Output for Section 2",
                    key="draft_summary_alerts_new_ades_followup",
                    height=320
                )

            # Section 3
            elif section["id"] == "actions_taken":
                regulatory_actions_file = st.file_uploader(
                    "Upload Regulatory Actions File",
                    type=["xlsx", "csv"],
                    key="regulatory_actions_upload"
                )

                if regulatory_actions_file is not None:
                    reg_df = read_uploaded_table(regulatory_actions_file)
                    if reg_df is not None:
                        st.session_state["regulatory_actions_df"] = reg_df
                        st.success("Regulatory actions file uploaded successfully.")
                        st.dataframe(reg_df.head(10), use_container_width=True)

                if st.button("Generate Section 3", key="btn_actions_taken"):
                    reg_df = st.session_state.get("regulatory_actions_df", None)

                    if reg_df is None:
                        draft_text = "No actions related to safety, labeling, or regulatory authority decisions were identified during the reporting interval."
                    else:
                        reg_summary = summarize_dataframe(reg_df, max_rows=8)
                        draft_text = generate_actions_taken_draft(
                            regulatory_actions_summary=reg_summary,
                            product_name=product_name,
                            interval_start=interval_start,
                            interval_end=interval_end
                        )

                    st.session_state["draft_actions_taken"] = draft_text

                st.text_area(
                    "Draft Output for 3. Actions Taken Since Last Periodic Adverse Drug Experience Report",
                    key="draft_actions_taken",
                    height=260
                )

            # Remaining standard sections
            else:
                source_data = st.text_area(
                    f"Source Data for {section['title']}",
                    key=f"source_{section['id']}",
                    height=180
                )

                comments = st.text_area(
                    f"Comments / Drafting Instructions for {section['title']}",
                    key=f"comment_{section['id']}",
                    height=100
                )

                if st.button(f"Generate Draft for {section['title']}", key=f"btn_{section['id']}"):
                    with st.spinner("Generating AI draft..."):
                        draft_text = generate_ai_draft(
                            section_title=section["title"],
                            section_purpose=section["purpose"],
                            source_data=source_data,
                            comments=comments,
                            product_name=product_name,
                            interval_start=interval_start,
                            interval_end=interval_end
                        )
                        st.session_state[f"draft_{section['id']}"] = draft_text

                st.text_area(
                    f"Draft Output for {section['title']}",
                    key=f"draft_{section['id']}",
                    height=220
                )

    # -----------------------------------------------------
    # Step 4: Assemble Full Report
    # -----------------------------------------------------
    st.header("Step 4: Assemble Full Report")

    if st.button("Assemble Full PADER Report"):
        full_report = assemble_full_report(
            product_name=product_name,
            interval_start=interval_start,
            interval_end=interval_end,
            report_owner=report_owner,
            sections=pader_sections
        )
        st.session_state["full_pader_report"] = full_report

    st.text_area(
        "Full PADER Report Output",
        key="full_pader_report",
        height=600
    )

    # -----------------------------------------------------
    # Step 5: Export
    # -----------------------------------------------------
    st.header("Step 5: Export")

    full_report_text = st.session_state.get("full_pader_report", "")
    if full_report_text:
        docx_bytes = export_report_to_word(full_report_text)
        st.download_button(
            label="Download Full PADER Report as Word",
            data=docx_bytes,
            file_name="PADER_Report.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
    else:
        st.info("Assemble the full report first to enable Word export.")

elif report_type in ["PBRER", "DSUR"]:
    st.info(f"{report_type} module coming soon.")
