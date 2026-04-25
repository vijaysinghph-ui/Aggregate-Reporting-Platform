import io

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

from utils.text import safe_text


REPORT_SECTION_HEADINGS = {
    "Cover Page",
    "Approval Page",
    "Table of Contents",
    "1. Introduction",
    "2. Summary of Submitted 15-Day Alerts, New Adverse Drug Experiences and New Adverse Drug Experience Follow-up",
    "3. Actions Taken Since Last Periodic Adverse Drug Experience Report",
    "4. Conclusion",
}


def set_document_style(doc: Document):
    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(10.5)

    for style_name in ["Heading 1", "Heading 2", "Title"]:
        if style_name in doc.styles:
            doc.styles[style_name].font.name = "Times New Roman"


def add_report_header_footer(doc: Document, full_report_text: str):
    lines = [line.strip() for line in full_report_text.splitlines() if line.strip()]
    title = "Annual Adverse Drug Experience Report"
    product = next((line for line in lines if line.startswith("Product:")), "")
    review_period = next((line for line in lines if line.startswith("Review Period:")), "")

    section = doc.sections[0]
    header = section.header
    header.paragraphs[0].text = title
    header.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    header.paragraphs[0].runs[0].bold = True

    if product or review_period:
        para = header.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.add_run(" | ".join([value for value in [product, review_period] if value]))

    footer = section.footer
    footer.paragraphs[0].text = "CONFIDENTIAL"
    footer.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.paragraphs[0].runs[0].bold = True


def add_centered_paragraph(doc: Document, text: str, bold: bool = False):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    run.bold = bold
    return paragraph


def parse_approval_page_text(text: str) -> tuple[list[str], list[dict]]:
    metadata = []
    signers = []
    current_signer = None

    for raw_line in text.splitlines():
        line = raw_line.strip()
        if not line:
            continue

        if line in {"Author:", "Medical Reviewer:", "Reviewed by:", "Approved by:"}:
            current_signer = {
                "role": line.rstrip(":"),
                "name": "",
                "designation": "",
                "organization": "",
            }
            signers.append(current_signer)
            continue

        if current_signer is None:
            metadata.append(line)
            continue

        if line.startswith("Name:"):
            current_signer["name"] = line.replace("Name:", "", 1).strip()
        elif line.startswith("Designation:"):
            current_signer["designation"] = line.replace("Designation:", "", 1).strip()
        elif line.startswith("For:"):
            current_signer["organization"] = line.replace("For:", "", 1).strip()

    return metadata, signers


def add_approval_signature_table(doc: Document, approval_text: str):
    metadata, signers = parse_approval_page_text(approval_text)

    for line in metadata:
        paragraph = doc.add_paragraph(line)
        paragraph.paragraph_format.space_after = Pt(2)

    doc.add_paragraph("")
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"

    headers = ["Role", "Name / Designation / Organization", "Signature / Date"]
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True

    for signer in signers:
        cells = table.add_row().cells
        cells[0].text = signer["role"]
        details = [
            signer.get("name", ""),
            signer.get("designation", ""),
            signer.get("organization", ""),
        ]
        cells[1].text = "\n".join([value for value in details if value])
        cells[2].text = "Signature: ____________________\nDate: ____________________"

    doc.add_paragraph("")


def parse_cover_page_text(text: str) -> dict:
    lines = [line.strip() for line in text.splitlines()]
    non_empty = [line for line in lines if line]

    metadata = {}
    confidentiality = ""
    company_lines = []
    descriptive_lines = []

    for line in non_empty:
        if line.startswith("Approval Date:"):
            metadata["Approval Date"] = line.replace("Approval Date:", "", 1).strip()
        elif line.startswith("Review Period:"):
            metadata["Review Period"] = line.replace("Review Period:", "", 1).strip()
        elif line.startswith("Report Status:"):
            metadata["Report Status"] = line.replace("Report Status:", "", 1).strip()
        elif line.startswith("Date of Report:"):
            metadata["Date of Report"] = line.replace("Date of Report:", "", 1).strip()
        elif line.startswith("This document is a confidential"):
            confidentiality = line
        elif line not in {
            "ANNUAL ADVERSE DRUG EXPERIENCE REPORT",
            "Periodic Adverse Drug Experience Report",
            "A report for the United States Food and Drug Administration",
        } and not line.startswith("(") and "NDA/ANDA No." not in line:
            if not metadata and not confidentiality:
                company_lines.append(line)
            else:
                descriptive_lines.append(line)

    title = non_empty[0] if len(non_empty) > 0 else "ANNUAL ADVERSE DRUG EXPERIENCE REPORT"
    period = non_empty[1] if len(non_empty) > 1 else ""
    product_line = non_empty[2] if len(non_empty) > 2 else ""
    product_display = product_line.split(";")[0].strip() if product_line else ""
    anda_number = ""
    if "NDA/ANDA No." in product_line:
        anda_number = product_line.split("NDA/ANDA No.", 1)[1].strip()

    if not confidentiality:
        confidentiality = " ".join(descriptive_lines)

    return {
        "title": title,
        "period": period,
        "product_line": product_line,
        "product_display": product_display,
        "anda_number": anda_number,
        "company_lines": company_lines,
        "metadata": metadata,
        "confidentiality": confidentiality,
    }


def add_cover_page_layout(doc: Document, cover_text: str):
    data = parse_cover_page_text(cover_text)

    add_centered_paragraph(doc, data["title"], bold=True)
    if data["period"]:
        add_centered_paragraph(doc, data["period"])
    if data["product_line"]:
        add_centered_paragraph(doc, data["product_line"], bold=True)

    doc.add_paragraph("")
    for line in data["company_lines"]:
        paragraph = doc.add_paragraph(line)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("")
    if data["product_display"]:
        add_centered_paragraph(doc, data["product_display"].upper(), bold=True)
    if data["anda_number"]:
        add_centered_paragraph(doc, f"ANDA Number: {data['anda_number']}")

    doc.add_paragraph("")
    add_centered_paragraph(doc, "Periodic Adverse Drug Experience Report", bold=True)
    add_centered_paragraph(doc, "A report for the United States Food and Drug Administration")

    doc.add_paragraph("")
    metadata_rows = [
        (label, value)
        for label, value in data["metadata"].items()
        if safe_text(value).strip()
    ]
    if metadata_rows:
        add_key_value_table(doc, metadata_rows)

    doc.add_paragraph("")
    if data["confidentiality"]:
        paragraph = doc.add_paragraph(data["confidentiality"])
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT


def parse_toc_text(text: str) -> list[str]:
    lines = []
    for raw_line in text.splitlines():
        line = raw_line.strip()
        if not line or line == "Cover Page":
            continue
        lines.append(line)
    return lines


def add_table_of_contents_layout(doc: Document, toc_text: str):
    entries = parse_toc_text(toc_text)
    add_centered_paragraph(doc, "Table of Contents", bold=True)

    if not entries:
        doc.add_paragraph("No table of contents entries available.")
        return

    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "Section"
    table.rows[0].cells[1].text = "Page"

    for idx, entry in enumerate(entries, start=1):
        cells = table.add_row().cells
        cells[0].text = entry
        cells[1].text = safe_text(idx)

    doc.add_paragraph("")
    doc.add_paragraph(
        "Note: Page numbers are placeholders in this prototype export and should be updated in Word before final submission."
    )


def is_markdown_table_row(line: str) -> bool:
    stripped = line.strip()
    return stripped.startswith("|") and stripped.endswith("|") and stripped.count("|") >= 2


def is_markdown_separator_row(line: str) -> bool:
    cells = parse_markdown_row(line)
    return bool(cells) and all(set(cell.replace(" ", "")) <= {"-"} for cell in cells)


def parse_markdown_row(line: str) -> list[str]:
    stripped = line.strip().strip("|")
    return [cell.strip() for cell in stripped.split("|")]


def add_markdown_table(doc: Document, table_lines: list[str]):
    if not table_lines:
        return

    rows = [parse_markdown_row(line) for line in table_lines if is_markdown_table_row(line)]
    rows = [row for row in rows if not all(set(cell.replace(" ", "")) <= {"-"} for cell in row)]
    if not rows:
        return

    max_cols = max(len(row) for row in rows)
    table = doc.add_table(rows=1, cols=max_cols)
    table.style = "Table Grid"

    for idx, value in enumerate(rows[0]):
        table.rows[0].cells[idx].text = value
        for paragraph in table.rows[0].cells[idx].paragraphs:
            for run in paragraph.runs:
                run.bold = True

    for row in rows[1:]:
        cells = table.add_row().cells
        for idx in range(max_cols):
            cells[idx].text = row[idx] if idx < len(row) else ""

    doc.add_paragraph("")


def add_lines_with_tables(doc: Document, lines: list[str]):
    table_buffer = []

    def flush_table():
        nonlocal table_buffer
        if table_buffer:
            add_markdown_table(doc, table_buffer)
            table_buffer = []

    for line in lines:
        stripped = line.strip()
        if is_markdown_table_row(stripped):
            table_buffer.append(stripped)
            continue

        flush_table()

        if not stripped:
            doc.add_paragraph("")
        elif stripped in {"=" * 80, "-" * 80}:
            continue
        else:
            paragraph = doc.add_paragraph(line)
            paragraph.paragraph_format.space_after = Pt(4)

    flush_table()


def generate_cover_page_text(
    product_name: str,
    dosage_strength: str,
    nda_anda_number: str,
    company_name: str,
    interval_start,
    interval_end,
    approval_date,
    report_status: str,
    report_status_other: str,
    report_date,
    confidentiality_statement: str,
    company_address: str,
) -> str:
    final_status = (
        report_status_other
        if report_status == "Other" and report_status_other
        else report_status
    )
    lines = [
        "ANNUAL ADVERSE DRUG EXPERIENCE REPORT",
        f"({interval_start} to {interval_end})",
        f"{product_name}, {dosage_strength}; NDA/ANDA No. {nda_anda_number}",
        "",
        f"{company_name}",
        company_address,
        "",
        "Periodic Adverse Drug Experience Report",
        "A report for the United States Food and Drug Administration",
        "",
        f"Approval Date: {approval_date}",
        f"Review Period: {interval_start} to {interval_end}",
        f"Report Status: {final_status}",
        f"Date of Report: {report_date}",
        "",
        confidentiality_statement,
    ]
    return "\n".join(lines)


def generate_approval_page_text(
    product_name: str,
    interval_start,
    interval_end,
    author_name: str,
    author_designation: str,
    medical_reviewer_name: str,
    medical_reviewer_designation: str,
    reviewer_name: str,
    reviewer_designation: str,
    approver_name: str,
    approver_designation: str,
    company_name: str,
) -> str:
    lines = [
        f"Product: {product_name}",
        f"Reporting Interval: {interval_start} to {interval_end}",
        "",
        "Author:",
        f"Name: {author_name}",
        f"Designation: {author_designation}",
        f"For: {company_name}",
        "Signature: ____________________",
        "Date: ____________________",
        "",
        "Medical Reviewer:",
        f"Name: {medical_reviewer_name}",
        f"Designation: {medical_reviewer_designation}",
        f"For: {company_name}",
        "Signature: ____________________",
        "Date: ____________________",
        "",
        "Reviewed by:",
        f"Name: {reviewer_name}",
        f"Designation: {reviewer_designation}",
        f"For: {company_name}",
        "Signature: ____________________",
        "Date: ____________________",
        "",
        "Approved by:",
        f"Name: {approver_name}",
        f"Designation: {approver_designation}",
        f"For: {company_name}",
        "Signature: ____________________",
        "Date: ____________________",
    ]
    return "\n".join(lines)


def generate_toc_text(sections: list[dict]) -> str:
    return "\n".join(
        section["title"] for section in sections if section["id"] != "cover_page"
    )


def assemble_full_report(
    product_name: str,
    interval_start,
    interval_end,
    report_owner: str,
    sections: list[dict],
    drafts: dict[str, str],
) -> str:
    report_parts = [
        "ANNUAL ADVERSE DRUG EXPERIENCE REPORT",
        f"Product: {product_name}",
        f"Review Period: {interval_start} to {interval_end}",
        f"Report Owner: {report_owner}",
        "\n" + "=" * 80 + "\n",
    ]

    for section in sections:
        section_title = section["title"]
        section_text = safe_text(drafts.get(section["id"], "")).strip()

        if not section_text:
            section_text = "[No draft available for this section yet.]"

        report_parts.append(section_title)
        report_parts.append(section_text)
        report_parts.append("\n" + "-" * 80 + "\n")

    return "\n".join(report_parts)


def export_report_to_word(full_report_text: str) -> bytes:
    doc = Document()
    set_document_style(doc)
    add_report_header_footer(doc, full_report_text)

    lines = full_report_text.splitlines()
    previous_was_section = False
    current_heading = None
    section_buffer = []

    def flush_section():
        nonlocal section_buffer
        if current_heading == "Cover Page":
            add_cover_page_layout(doc, "\n".join(section_buffer))
        elif current_heading == "Approval Page":
            add_approval_signature_table(doc, "\n".join(section_buffer))
        elif current_heading == "Table of Contents":
            add_table_of_contents_layout(doc, "\n".join(section_buffer))
        else:
            add_lines_with_tables(doc, section_buffer)
        section_buffer = []

    for line in lines:
        stripped = line.strip()

        if stripped == "ANNUAL ADVERSE DRUG EXPERIENCE REPORT":
            flush_section()
            add_centered_paragraph(doc, stripped, bold=True)
        elif stripped in REPORT_SECTION_HEADINGS:
            flush_section()
            if previous_was_section is False and stripped != "Cover Page":
                doc.add_page_break()
            doc.add_heading(stripped, level=1)
            current_heading = stripped
            previous_was_section = True
        else:
            section_buffer.append(line)
            previous_was_section = False

    flush_section()

    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio.getvalue()


def add_key_value_table(doc: Document, values: list[tuple[str, object]]):
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "Field"
    table.rows[0].cells[1].text = "Value"

    for label, value in values:
        cells = table.add_row().cells
        cells[0].text = safe_text(label)
        cells[1].text = safe_text(value)


def add_records_table(doc: Document, headers: list[str], records: list[dict]):
    if not records:
        doc.add_paragraph("No records available.")
        return

    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for idx, header in enumerate(headers):
        table.rows[0].cells[idx].text = header

    for record in records:
        cells = table.add_row().cells
        for idx, header in enumerate(headers):
            cells[idx].text = safe_text(record.get(header, ""))


def export_audit_trail_to_word(
    report_title: str,
    context: dict,
    approval_context: dict,
    workflow_status: str,
    workflow_history: list[dict],
    review_comments: list[dict],
    versions: list[dict],
) -> bytes:
    doc = Document()
    doc.add_heading("PADER Audit Trail", level=0)

    doc.add_heading("Report Summary", level=1)
    add_key_value_table(
        doc,
        [
            ("Report Title", report_title),
            ("Product Name", context.get("product_name", "")),
            ("NDA / ANDA Number", context.get("nda_anda_number", "")),
            ("Company Name", context.get("company_name", "")),
            ("Reporting Interval Start", context.get("interval_start", "")),
            ("Reporting Interval End", context.get("interval_end", "")),
            ("Current Workflow Status", workflow_status),
        ],
    )

    doc.add_heading("Assigned Roles", level=1)
    add_key_value_table(
        doc,
        [
            ("Author", approval_context.get("author_name", "")),
            ("Author Designation", approval_context.get("author_designation", "")),
            ("Medical Reviewer", approval_context.get("medical_reviewer_name", "")),
            ("Reviewer", approval_context.get("reviewer_name", "")),
            ("Reviewer Designation", approval_context.get("reviewer_designation", "")),
            ("Approver", approval_context.get("approver_name", "")),
            ("Approver Designation", approval_context.get("approver_designation", "")),
        ],
    )

    doc.add_heading("Workflow History", level=1)
    add_records_table(
        doc,
        ["timestamp", "role", "actor", "action", "comment"],
        workflow_history,
    )

    doc.add_heading("Review Comments", level=1)
    add_records_table(
        doc,
        ["timestamp", "role", "actor", "comment"],
        review_comments,
    )

    doc.add_heading("Version History", level=1)
    version_records = [
        {
            "version": version.get("version_label", ""),
            "type": version.get("version_type", ""),
            "timestamp": version.get("timestamp", ""),
            "role": version.get("actor_role", ""),
            "actor": version.get("actor_name", ""),
            "action": version.get("action", ""),
            "status": version.get("workflow_status", ""),
        }
        for version in versions
    ]
    add_records_table(
        doc,
        ["version", "type", "timestamp", "role", "actor", "action", "status"],
        version_records,
    )

    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio.getvalue()
